import os
import requests
import base64
from docx import Document
from langchain_ollama import OllamaLLM
from langchain_core.messages import HumanMessage
from rich.console import Console
from rich.live import Live
from rich.text import Text

console = Console()

def fetch_github_pr(url, token=None):
    try:
        parts = [p for p in url.rstrip('/').split('/') if p]
        
        if len(parts) >= 4 and parts[-2] == "pull":
            owner, repo, pr_number = parts[-4], parts[-3], parts[-1]
            api_url = f"https://api.github.com/repos/{owner}/{repo}/pulls/{pr_number}"
        elif len(parts) >= 2:
            owner, repo = parts[-2], parts[-1]
            console.print(f"[bold yellow]Note: You provided a repository link, not a specific PR.\nWe will fetch the most recently updated PR for {owner}/{repo}...[/bold yellow]")
            search_url = f"https://api.github.com/repos/{owner}/{repo}/pulls?state=all&sort=updated&direction=desc&per_page=1"
            
            search_headers = {'Accept': 'application/vnd.github.v3+json'}
            if token:
                search_headers['Authorization'] = f'token {token}'
                
            search_res = requests.get(search_url, headers=search_headers)
            if search_res.status_code == 200 and search_res.json():
                pr_number = search_res.json()[0]['number']
                api_url = f"https://api.github.com/repos/{owner}/{repo}/pulls/{pr_number}"
            else:
                console.print(f"[bold red]Could not find any Pull Requests for this repository.[/bold red]")
                return None
        else:
            console.print(f"[bold red]Invalid GitHub URL format.[/bold red]")
            return None
            
        headers = {'Accept': 'application/vnd.github.v3.diff'}
        if token:
            headers['Authorization'] = f'token {token}'
            
        diff_url = f"{api_url}.diff"
        
        console.print(f"[dim]Fetching PR metadata from {api_url}...[/dim]")
        
        meta_headers = {'Accept': 'application/vnd.github.v3+json'}
        if token:
            meta_headers['Authorization'] = f'token {token}'
        meta_res = requests.get(api_url, headers=meta_headers)
        
        if meta_res.status_code != 200:
            console.print(f"[bold red]Error fetching metadata: HTTP {meta_res.status_code}[/bold red]")
            return None
            
        metadata = meta_res.json()
        title = metadata.get('title', 'Unknown Title')
        body = metadata.get('body', 'No description provided.')
        
        console.print(f"[dim]Fetching PR diff...[/dim]")
        diff_res = requests.get(diff_url, headers=headers)
        if diff_res.status_code != 200:
            console.print(f"[bold red]Error fetching diff: HTTP {diff_res.status_code}[/bold red]")
            return None
            
        diff_text = diff_res.text
        return {
            "title": title,
            "body": body,
            "url": url,
            "diff": diff_text
        }
    except Exception as e:
        console.print(f"[bold red]Error parsing URL or fetching data: {e}[/bold red]")
        return None

def save_to_word(text, title, output_dir):
    try:
        doc = Document()
        doc.add_heading(f"PR Summary: {title}", level=0)
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Remove any leading markdown bullet asterisks or dashes for list formatting
            is_list = False
            if line.startswith('- ') or line.startswith('* '):
                is_list = True
                line = line[2:].strip()
                
            # Handle Headings
            if line.startswith('# '):
                doc.add_heading(line[2:].replace('**', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].replace('**', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].replace('**', ''), level=3)
            else:
                # Add a normal or list paragraph
                p = doc.add_paragraph(style='List Bullet' if is_list else 'Normal')
                
                # Parse inline bolding (**text**)
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    # Every odd index in a ** split means it was inside asterisks (assuming balanced pairs)
                    if i % 2 != 0:
                        run.bold = True
        
        os.makedirs(output_dir, exist_ok=True)
        safe_title = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        safe_title = safe_title.replace(" ", "_")
        if not safe_title:
            safe_title = "PR_Summary"
            
        file_path = os.path.join(output_dir, f"{safe_title}.docx")
        doc.save(file_path)
        console.print(f"\n[bold green]Success! Document saved to:[/bold green] {file_path}\n")
    except Exception as e:
        console.print(f"[bold red]Error saving Word Document: {e}[/bold red]")

def generate_doc_from_pr(pr_data, output_dir):
    console.print(f"\n[cyan]Generating Documentation with qwen2.5:14b...[/cyan]")
    try:
        llm = OllamaLLM(model="qwen2.5:14b", num_thread=8)
        
        prompt = f"""You are a Senior Technical Writer. Analyze the following Pull Request details and generate a structured business document summary.
Do not output raw Markdown markdown blocks (```) covering the entire response, just output the formatted text.
Create headers for "Executive Summary", "Technical Breakdown", and "Potential Impacts".

Pull Request Title: {pr_data['title']}
Description: {pr_data['body']}

Git Diff (Truncated):
{pr_data['diff'][:15000]}
"""
        response = ""
        with Live(Text("...", style="blink dim"), refresh_per_second=10, transient=True) as live:
            for chunk in llm.stream(prompt):
                response += chunk
                live.update(Text("Generating... " + str(len(response)) + " chars read"))
        
        save_to_word(response, pr_data['title'], output_dir)
    except Exception as e:
        console.print(f"[bold red]Model inference error (ensure qwen2.5:14b is pulled!): {e}[/bold red]")

def generate_doc_from_image(image_path, output_dir):
    console.print(f"\n[cyan]Analyzing image with llama3.2-vision:11b...[/cyan]")
    
    if not os.path.exists(image_path):
        console.print("[bold red]Image path not found.[/bold red]")
        return
        
    try:
        with open(image_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode('utf-8')

        llm = OllamaLLM(model="llama3.2-vision:11b", num_thread=8)
        
        message = HumanMessage(
            content=[
                {"type": "text", "text": "You are a Technical Writer. Look at this screenshot of code changes or PR diff, and generate a structured business document summarizing the changes. Include an 'Executive Summary', 'Technical Breakdown', and 'Potential Impacts'. Do not wrap the entire response in markdown code blocks."},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encoded_string}"}}
            ]
        )
        
        response = ""
        with Live(Text("...", style="blink dim"), refresh_per_second=10, transient=True) as live:
            for chunk in llm.stream([message]):
                response += chunk
                live.update(Text("Processing Image... " + str(len(response)) + " chars read"))
        
        title = "Screenshot_Summary"
        save_to_word(response, title, output_dir)
    except Exception as e:
        console.print(f"[bold red]Vision processing error (make sure llama3.2-vision:11b is downloaded and running!): {e}[/bold red]")

def run(data_dir):
    docs_dir = os.path.join(data_dir, "docs")
    console.print("\n[bold cyan]─── PR Documentation Generator ───[/bold cyan]")
    console.print("[dim]Enter a GitHub Pull Request URL (e.g. https://github.com/org/repo/pull/1) \nOR an absolute path to a screenshot image (e.g. /Users/name/Desktop/pr.png)[/dim]")
    
    user_input = console.input("\n[bold orange3]URL or Image Path:[/bold orange3] ").strip()
    
    # Clean up drag-and-drop path artifacts (Quotes or terminal escapes like \ and \()
    user_input = user_input.replace("'", "").replace('"', "").strip()
    user_input = user_input.replace("\\", "")
    
    if not user_input:
        return
        
    if user_input.startswith("http"):
        token = console.input("[dim]GitHub PAT Token (Optional for public repos, press Enter to skip): [/dim]").strip()
        pr_data = fetch_github_pr(user_input, token if token else None)
        if pr_data:
            generate_doc_from_pr(pr_data, docs_dir)
    elif os.path.exists(user_input):
        generate_doc_from_image(user_input, docs_dir)
    else:
        console.print(f"[bold red]Invalid input. Must be a valid HTTP URL or an existing file path. Received: {user_input}[/bold red]")
