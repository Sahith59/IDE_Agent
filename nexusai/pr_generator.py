"""PR Documentation Generator — adapted from pr_generator/generator.py"""
import os
import requests
import base64
from docx import Document
from langchain_ollama import OllamaLLM
from langchain_core.messages import HumanMessage
from rich.console import Console
from rich.live import Live
from rich.text import Text

console = Console()


def fetch_github_pr(url, token=None):
    try:
        parts = [p for p in url.rstrip('/').split('/') if p]

        if len(parts) >= 4 and parts[-2] == "pull":
            owner, repo, pr_number = parts[-4], parts[-3], parts[-1]
            api_url = f"https://api.github.com/repos/{owner}/{repo}/pulls/{pr_number}"
        elif len(parts) >= 2:
            owner, repo = parts[-2], parts[-1]
            console.print(f"[bold yellow]Fetching most recently updated PR for {owner}/{repo}...[/bold yellow]")
            search_url  = f"https://api.github.com/repos/{owner}/{repo}/pulls?state=all&sort=updated&direction=desc&per_page=1"
            h = {'Accept': 'application/vnd.github.v3+json'}
            if token: h['Authorization'] = f'token {token}'
            res = requests.get(search_url, headers=h)
            if res.status_code == 200 and res.json():
                pr_number = res.json()[0]['number']
                api_url   = f"https://api.github.com/repos/{owner}/{repo}/pulls/{pr_number}"
            else:
                console.print("[bold red]Could not find any Pull Requests.[/bold red]")
                return None
        else:
            console.print("[bold red]Invalid GitHub URL format.[/bold red]")
            return None

        meta_h = {'Accept': 'application/vnd.github.v3+json'}
        if token: meta_h['Authorization'] = f'token {token}'

        console.print(f"[dim]Fetching PR metadata...[/dim]")
        meta_res = requests.get(api_url, headers=meta_h)
        if meta_res.status_code != 200:
            console.print(f"[bold red]Error fetching metadata: HTTP {meta_res.status_code}[/bold red]")
            return None

        metadata = meta_res.json()

        diff_h = {'Accept': 'application/vnd.github.v3.diff'}
        if token: diff_h['Authorization'] = f'token {token}'

        console.print("[dim]Fetching PR diff...[/dim]")
        diff_res = requests.get(f"{api_url}.diff", headers=diff_h)
        if diff_res.status_code != 200:
            console.print(f"[bold red]Error fetching diff: HTTP {diff_res.status_code}[/bold red]")
            return None

        return {
            "title": metadata.get('title', 'Unknown Title'),
            "body":  metadata.get('body', 'No description provided.'),
            "url":   url,
            "diff":  diff_res.text,
        }
    except Exception as e:
        console.print(f"[bold red]Error: {e}[/bold red]")
        return None


def save_to_word(text, title, output_dir):
    try:
        doc = Document()
        doc.add_heading(f"PR Summary: {title}", level=0)

        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            is_list = False
            if line.startswith('- ') or line.startswith('* '):
                is_list = True
                line    = line[2:].strip()

            if line.startswith('# '):
                doc.add_heading(line[2:].replace('**', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].replace('**', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].replace('**', ''), level=3)
            else:
                p     = doc.add_paragraph(style='List Bullet' if is_list else 'Normal')
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run      = p.add_run(part)
                    run.bold = (i % 2 != 0)

        os.makedirs(output_dir, exist_ok=True)
        safe = "".join(c for c in title if c.isalpha() or c.isdigit() or c == ' ').rstrip()
        safe = safe.replace(" ", "_") or "PR_Summary"
        path = os.path.join(output_dir, f"{safe}.docx")
        doc.save(path)
        console.print(f"\n[bold green]Saved:[/bold green] {path}\n")
    except Exception as e:
        console.print(f"[bold red]Error saving document: {e}[/bold red]")


def generate_doc_from_pr(pr_data, output_dir):
    console.print("\n[cyan]Generating documentation with qwen2.5:14b...[/cyan]")
    try:
        llm    = OllamaLLM(model="qwen2.5:14b", num_thread=8)
        prompt = (
            "You are a Senior Technical Writer. Analyze the following Pull Request and generate "
            "a structured business document summary. Do not wrap the entire response in markdown "
            "code blocks. Create headers for 'Executive Summary', 'Technical Breakdown', and "
            "'Potential Impacts'.\n\n"
            f"Pull Request Title: {pr_data['title']}\n"
            f"Description: {pr_data['body']}\n\n"
            f"Git Diff (Truncated):\n{pr_data['diff'][:15000]}"
        )
        response = ""
        with Live(Text("...", style="blink dim"), refresh_per_second=10, transient=True) as live:
            for chunk in llm.stream(prompt):
                response += chunk
                live.update(Text(f"Generating... {len(response)} chars"))
        save_to_word(response, pr_data['title'], output_dir)
    except Exception as e:
        console.print(f"[bold red]Model error: {e}[/bold red]")


def generate_doc_from_image(image_path, output_dir):
    console.print("\n[cyan]Analyzing image with llama3.2-vision:11b...[/cyan]")
    if not os.path.exists(image_path):
        console.print("[bold red]Image path not found.[/bold red]")
        return
    try:
        with open(image_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")
        llm     = OllamaLLM(model="llama3.2-vision:11b", num_thread=8)
        message = HumanMessage(content=[
            {"type": "text", "text": "You are a Technical Writer. Look at this screenshot of code changes or PR diff, and generate a structured business document summarizing the changes. Include an 'Executive Summary', 'Technical Breakdown', and 'Potential Impacts'. Do not wrap the entire response in markdown code blocks."},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encoded}"}}
        ])
        response = ""
        with Live(Text("...", style="blink dim"), refresh_per_second=10, transient=True) as live:
            for chunk in llm.stream([message]):
                response += chunk
                live.update(Text(f"Processing... {len(response)} chars"))
        save_to_word(response, "Screenshot_Summary", output_dir)
    except Exception as e:
        console.print(f"[bold red]Vision error: {e}[/bold red]")


def run(nexus_home: str):
    docs_dir = os.path.join(nexus_home, "docs")
    console.print("\n[bold cyan]─── PR Documentation Generator ───[/bold cyan]")
    console.print("[dim]Enter a GitHub PR URL or an absolute path to a screenshot image.[/dim]")

    user_input = console.input("\n[bold orange3]URL or Image Path:[/bold orange3] ").strip()
    user_input = user_input.replace("'", "").replace('"', "").replace("\\", "").strip()

    if not user_input:
        return

    if user_input.startswith("http"):
        token   = console.input("[dim]GitHub PAT Token (optional, Enter to skip): [/dim]").strip()
        pr_data = fetch_github_pr(user_input, token or None)
        if pr_data:
            generate_doc_from_pr(pr_data, docs_dir)
    elif os.path.exists(user_input):
        generate_doc_from_image(user_input, docs_dir)
    else:
        console.print(f"[bold red]Invalid input: must be a URL or existing file path.[/bold red]")
